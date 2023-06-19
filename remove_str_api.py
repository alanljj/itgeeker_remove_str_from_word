# -*- coding: utf-8 -*-
###########################################################################
#    Copyright 2023 奇客罗方智能科技 https://www.geekercloud.com
#    ITGeeker.net <alanljj@gmail.com>
############################################################################
import glob
import os
import re
from docx import Document
from remove_str_convert_doc2docx import convert_doc2docx_by_win32com


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            ptext = regex.sub(replace, p.text)
            p.text = ptext
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def generate_file_and_str_list(report_p, val_list):
    revised_p = os.path.join(report_p, '已处理文件')
    print('revised_p: ', revised_p)
    if not os.path.exists(revised_p):
        os.makedirs(revised_p)
    docx_list = [f for f in glob.glob(report_p + r"\[!~$]*.docx")]
    print('docx_list: ', docx_list)
    replaced = replace_str_for_file_list(docx_list, val_list, revised_p)
    if replaced:
        return len(docx_list)
    return False

def replace_str_for_file_list(docx_list, val_list, revised_p):
    for docx_f in docx_list:
        doc = Document(docx_f)
        filename, file_extension = os.path.splitext(docx_f)
        basename_no_ext = os.path.splitext(os.path.basename(docx_f))[0]
        # print('filename: ', filename)
        print('file_extension: ', file_extension)
        print('basename: ', basename_no_ext)

        revised_file = os.path.join(revised_p, basename_no_ext + '-revised' + file_extension)
        print('revised_file: ', revised_file)

        for val in val_list:
            # re_str = "\b" + val + "\b"
            # regex1 = re.compile(val)
            # regex1 = re.compile(r'(?i)\b%s\b' % val)
            regex1 = re.compile(r'\b{}\b'.format(val), flags=re.I)
            #         # regex2 = re.compile('.*({}).*'.format(what2look4))
            # print('regex1: ', regex1)

            replace1 = r""
            docx_replace_regex(doc, regex1, replace1)
        doc.save(revised_file)
    return True
